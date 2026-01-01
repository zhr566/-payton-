# 毕业论文格式大师.exe
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os
import threading
import time


class ThesisFormatter:
    def __init__(self, root):
        self.root = root
        self.root.title("蚌埠医科大学毕业论文格式大师")
        self.root.geometry("800x600")
        self.root.configure(bg='#f0f0f0')

        # 设置图标（如果有的话）
        try:
            self.root.iconbitmap('icon.ico')
        except:
            pass

        self.setup_ui()

    def setup_ui(self):
        # 标题栏
        title_frame = tk.Frame(self.root, bg='#2c3e50', height=80)
        title_frame.pack(fill='x', side='top')
        title_frame.pack_propagate(False)

        title_label = tk.Label(title_frame, text="毕业论文格式大师",
                               font=('微软雅黑', 24, 'bold'),
                               fg='white', bg='#2c3e50')
        title_label.pack(expand=True)

        subtitle_label = tk.Label(title_frame, text="蚌埠医科大学专用版",
                                  font=('微软雅黑', 12),
                                  fg='#ecf0f1', bg='#2c3e50')
        subtitle_label.pack()

        # 主内容区
        main_frame = tk.Frame(self.root, bg='#f0f0f0')
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # 左侧功能区
        left_frame = tk.Frame(main_frame, bg='#f0f0f0')
        left_frame.pack(side='left', fill='y', padx=(0, 20))

        # 文件选择区域
        file_frame = tk.LabelFrame(left_frame, text="📁 文件选择",
                                   font=('微软雅黑', 12, 'bold'),
                                   bg='#ffffff', fg='#2c3e50',
                                   relief='groove', bd=2)
        file_frame.pack(fill='x', pady=(0, 20))

        # 文件路径显示
        self.file_path_var = tk.StringVar()
        path_entry = tk.Entry(file_frame, textvariable=self.file_path_var,
                              font=('微软雅黑', 10), width=30,
                              relief='solid', bd=1)
        path_entry.pack(side='left', padx=10, pady=15, fill='x', expand=True)

        # 浏览按钮
        browse_btn = tk.Button(file_frame, text="浏览...",
                               command=self.browse_file,
                               font=('微软雅黑', 10, 'bold'),
                               bg='#3498db', fg='white',
                               activebackground='#2980b9',
                               relief='raised', bd=2,
                               cursor='hand2')
        browse_btn.pack(side='right', padx=(0, 10), pady=15)

        # 格式设置区域
        format_frame = tk.LabelFrame(left_frame, text="⚙️ 格式设置",
                                     font=('微软雅黑', 12, 'bold'),
                                     bg='#ffffff', fg='#2c3e50',
                                     relief='groove', bd=2)
        format_frame.pack(fill='x', pady=(0, 20))

        # 格式选项
        self.font_chinese_var = tk.StringVar(value="宋体")
        self.font_english_var = tk.StringVar(value="Times New Roman")
        self.font_size_var = tk.StringVar(value="10.5")
        self.line_spacing_var = tk.StringVar(value="1.5")
        self.margin_var = tk.StringVar(value="2.5")

        options = [
            ("中文字体:", self.font_chinese_var),
            ("英文字体:", self.font_english_var),
            ("字号(磅):", self.font_size_var),
            ("行距(倍):", self.line_spacing_var),
            ("页边距(cm):", self.margin_var)
        ]

        for i, (label, var) in enumerate(options):
            frame = tk.Frame(format_frame, bg='white')
            frame.pack(fill='x', padx=10, pady=5)

            tk.Label(frame, text=label, font=('微软雅黑', 10),
                     bg='white', width=10, anchor='w').pack(side='left')

            entry = tk.Entry(frame, textvariable=var, font=('微软雅黑', 10),
                             width=15, relief='solid', bd=1)
            entry.pack(side='right')

        # 处理按钮
        process_btn = tk.Button(left_frame, text="✨ 开始格式化",
                                command=self.start_formatting,
                                font=('微软雅黑', 14, 'bold'),
                                bg='#27ae60', fg='white',
                                activebackground='#229954',
                                relief='raised', bd=3,
                                cursor='hand2',
                                height=2)
        process_btn.pack(fill='x', pady=10)

        # 右侧日志区域
        right_frame = tk.Frame(main_frame, bg='#f0f0f0')
        right_frame.pack(side='right', fill='both', expand=True)

        log_frame = tk.LabelFrame(right_frame, text="📝 处理日志",
                                  font=('微软雅黑', 12, 'bold'),
                                  bg='#ffffff', fg='#2c3e50',
                                  relief='groove', bd=2)
        log_frame.pack(fill='both', expand=True)

        # 日志文本框
        self.log_text = scrolledtext.ScrolledText(log_frame,
                                                  font=('Consolas', 10),
                                                  bg='#2c3e50', fg='#ecf0f1',
                                                  insertbackground='white',
                                                  relief='flat')
        self.log_text.pack(fill='both', expand=True, padx=5, pady=5)

        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(right_frame,
                                            variable=self.progress_var,
                                            maximum=100,
                                            mode='determinate',
                                            length=100)
        self.progress_bar.pack(fill='x', pady=(10, 0))

        # 状态栏
        status_frame = tk.Frame(self.root, bg='#34495e', height=30)
        status_frame.pack(fill='x', side='bottom')
        status_frame.pack_propagate(False)

        self.status_label = tk.Label(status_frame, text="就绪",
                                     font=('微软雅黑', 9),
                                     fg='white', bg='#34495e',
                                     anchor='w')
        self.status_label.pack(side='left', padx=10)

        # 版本信息
        version_label = tk.Label(status_frame, text="v1.0 © 2025",
                                 font=('微软雅黑', 9),
                                 fg='#bdc3c7', bg='#34495e',
                                 anchor='e')
        version_label.pack(side='right', padx=10)

    def browse_file(self):
        file_path = filedialog.askopenfilename(
            title="选择毕业论文文件",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if file_path:
            self.file_path_var.set(file_path)
            self.log(f"已选择文件: {os.path.basename(file_path)}")

    def log(self, message):
        timestamp = time.strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update()

    def update_status(self, message):
        self.status_label.config(text=message)
        self.root.update()

    def update_progress(self, value):
        self.progress_var.set(value)
        self.root.update()

    def start_formatting(self):
        file_path = self.file_path_var.get()

        if not file_path:
            messagebox.showwarning("警告", "请先选择要处理的文件！")
            return

        if not os.path.exists(file_path):
            messagebox.showerror("错误", f"文件不存在:\n{file_path}")
            return

        # 在新线程中处理，避免界面卡顿
        thread = threading.Thread(target=self.format_document, args=(file_path,))
        thread.daemon = True
        thread.start()

    def format_document(self, input_path):
        try:
            self.update_status("正在处理...")
            self.update_progress(10)
            self.log("=" * 50)
            self.log("开始格式化论文")
            self.log("=" * 50)

            # 读取文档
            self.log(f"📖 读取文档: {os.path.basename(input_path)}")
            doc = Document(input_path)
            self.update_progress(30)

            # 设置页面边距
            margin_cm = float(self.margin_var.get())
            self.log(f"📐 设置页边距: {margin_cm}cm")
            for section in doc.sections:
                section.top_margin = Cm(margin_cm)
                section.bottom_margin = Cm(margin_cm)
                section.left_margin = Cm(margin_cm)
                section.right_margin = Cm(margin_cm)
            self.update_progress(40)

            # 获取设置值
            font_chinese = self.font_chinese_var.get()
            font_english = self.font_english_var.get()
            font_size = float(self.font_size_var.get())
            line_spacing = float(self.line_spacing_var.get())

            # 格式化所有段落
            self.log("🎨 应用字体和段落格式...")
            para_count = 0
            run_count = 0

            for i, para in enumerate(doc.paragraphs):
                # 设置行距
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(font_size * line_spacing)

                # 正文段落设置首行缩进
                if len(para.text.strip()) > 10:
                    para.paragraph_format.first_line_indent = Cm(0.74)

                # 处理每个文字片段
                for run in para.runs:
                    if run.text.strip():
                        # 设置字号
                        run.font.size = Pt(font_size)

                        # 判断中英文并设置字体
                        text = run.text
                        if any('\u4e00' <= c <= '\u9fff' for c in text):
                            # 中文
                            run.font.name = font_chinese
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_chinese)
                        else:
                            # 英文
                            run.font.name = font_english

                        run_count += 1

                para_count += 1

                # 更新进度
                if i % 10 == 0:
                    progress = 40 + (i / len(doc.paragraphs)) * 40
                    self.update_progress(progress)

            self.update_progress(80)
            self.log(f"✓ 已处理 {para_count} 个段落，{run_count} 个文字片段")

            # 处理表格
            self.log("📊 处理表格格式...")
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text.strip():
                                    run.font.size = Pt(font_size)
                table_count += 1
            self.log(f"✓ 已处理 {table_count} 个表格")

            self.update_progress(90)

            # 保存新文件
            dir_name = os.path.dirname(input_path)
            base_name = os.path.basename(input_path)
            new_name = base_name.replace('.docx', '_格式正确.docx')
            output_path = os.path.join(dir_name, new_name)

            # 防止重名
            counter = 1
            while os.path.exists(output_path):
                new_name = base_name.replace('.docx', f'_格式正确({counter}).docx')
                output_path = os.path.join(dir_name, new_name)
                counter += 1

            self.log(f"💾 保存文件: {new_name}")
            doc.save(output_path)

            self.update_progress(100)

            # 显示完成信息
            self.log("=" * 50)
            self.log("✅ 格式化完成！")
            self.log(f"📄 新文件: {output_path}")
            self.log("=" * 50)

            self.update_status("处理完成")

            # 弹出成功对话框
            self.root.after(0, lambda: messagebox.showinfo(
                "完成",
                f"论文格式化完成！\n\n"
                f"原文件: {base_name}\n"
                f"新文件: {new_name}\n\n"
                f"已保存到相同目录。"
            ))

        except Exception as e:
            self.log(f"❌ 错误: {str(e)}")
            self.update_status("处理失败")
            self.root.after(0, lambda: messagebox.showerror("错误", f"处理失败:\n{str(e)}"))


def main():
    root = tk.Tk()
    app = ThesisFormatter(root)
    root.mainloop()


if __name__ == "__main__":
    main()